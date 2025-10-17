"""File content extraction provider implementation."""
import os
from pathlib import Path

from ...core.interfaces.providers.i_content_extraction_provider import IContentExtractionProvider
from ...core.results.result import Result
from ...core.value_objects.enums import FileType


class FileContentExtractionProvider(IContentExtractionProvider):
    """
    Concrete implementation of IContentExtractionProvider.

    Extracts text content from various file formats using appropriate libraries.
    Supports PDF, DOCX, DOC, TXT, and Markdown files.
    """

    def extract_text_from_pdf(self, file_path: str) -> Result[str]:
        """
        Extract text content from a PDF file using PyPDF2.

        Args:
            file_path: Path to the PDF file

        Returns:
            Result[str]: Success with extracted text or failure
        """
        try:
            import PyPDF2
        except ImportError:
            return Result.failure("PyPDF2 library not installed. Install with: pip install PyPDF2")

        if not os.path.exists(file_path):
            return Result.failure(f"File not found: {file_path}")

        try:
            text_content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            extracted_text = "\n\n".join(text_content)

            if not extracted_text.strip():
                return Result.failure("No text content could be extracted from PDF (may be image-based)")

            return Result.success(extracted_text)

        except Exception as e:
            return Result.failure(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> Result[str]:
        """
        Extract text content from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Result[str]: Success with extracted text or failure
        """
        try:
            import docx
        except ImportError:
            return Result.failure("python-docx library not installed. Install with: pip install python-docx")

        if not os.path.exists(file_path):
            return Result.failure(f"File not found: {file_path}")

        try:
            doc = docx.Document(file_path)

            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            extracted_text = "\n\n".join(text_content)

            if not extracted_text.strip():
                return Result.failure("No text content could be extracted from DOCX")

            return Result.success(extracted_text)

        except Exception as e:
            return Result.failure(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_doc(self, file_path: str) -> Result[str]:
        """
        Extract text content from a DOC file using antiword subprocess.

        Args:
            file_path: Path to the DOC file

        Returns:
            Result[str]: Success with extracted text or failure
        """
        if not os.path.exists(file_path):
            return Result.failure(f"File not found: {file_path}")

        # Try using antiword via subprocess (available on most Linux systems)
        try:
            import subprocess

            # Try antiword command (if installed on the system)
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode == 0:
                extracted_text = result.stdout

                if not extracted_text.strip():
                    return Result.failure("No text content could be extracted from DOC")

                return Result.success(extracted_text)
            else:
                return Result.failure(
                    f"antiword command failed with return code {result.returncode}. "
                    f"Error: {result.stderr}\n"
                    f"Install antiword: sudo apt-get install antiword (Linux) or brew install antiword (macOS)"
                )

        except FileNotFoundError:
            return Result.failure(
                "antiword command not found. DOC file extraction requires antiword.\n"
                "Install antiword:\n"
                "  - Linux: sudo apt-get install antiword\n"
                "  - macOS: brew install antiword\n"
                "  - Windows: Use LibreOffice or convert to DOCX format"
            )
        except subprocess.TimeoutExpired:
            return Result.failure("DOC extraction timed out (exceeded 30 seconds)")
        except Exception as e:
            return Result.failure(f"Failed to extract text from DOC: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> Result[str]:
        """
        Extract text content from a plain text file.

        Args:
            file_path: Path to the text file

        Returns:
            Result[str]: Success with extracted text or failure
        """
        if not os.path.exists(file_path):
            return Result.failure(f"File not found: {file_path}")

        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()

                        if not text.strip():
                            return Result.failure("Text file is empty")

                        return Result.success(text)
                except UnicodeDecodeError:
                    continue

            return Result.failure(f"Failed to decode text file with any supported encoding: {encodings}")

        except Exception as e:
            return Result.failure(f"Failed to read text file: {str(e)}")

    def extract_text_from_markdown(self, file_path: str) -> Result[str]:
        """
        Extract text content from a Markdown file.

        Markdown files are plain text, so this is similar to txt extraction.

        Args:
            file_path: Path to the Markdown file

        Returns:
            Result[str]: Success with extracted text or failure
        """
        if not os.path.exists(file_path):
            return Result.failure(f"File not found: {file_path}")

        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()

                        if not text.strip():
                            return Result.failure("Markdown file is empty")

                        return Result.success(text)
                except UnicodeDecodeError:
                    continue

            return Result.failure(f"Failed to decode Markdown file with any supported encoding: {encodings}")

        except Exception as e:
            return Result.failure(f"Failed to read Markdown file: {str(e)}")
